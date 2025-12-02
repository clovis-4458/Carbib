import pandas as pd
from datetime import date, timedelta

from django.http import HttpResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth import authenticate, login
from django.contrib.auth.decorators import login_required
from django.views.decorators.cache import never_cache, cache_control
from django.contrib.auth.views import LogoutView
from django.template.loader import get_template

from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches

from .forms import CustomAuthenticationForm, RegistrationForm, CandidateApplicationForm
from .models import Candidates, Jobs, Agents, Countries


# ----------------------------- HOME / REGISTRATION -----------------------------
def home(request):
    if request.method == "POST":
        form = RegistrationForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Registration successful! You can now log in.")
            return redirect("login")
        else:
            messages.error(request, "Please correct the errors below.")
    else:
        form = RegistrationForm()
    return render(request, "myapp/index.html", {"form": form})


# ----------------------------------- LOGIN ------------------------------------
def login_view(request):
    if request.user.is_authenticated:
        return redirect("dashboard_view")

    if request.method == "POST":
        form = CustomAuthenticationForm(request, data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get("username")
            password = form.cleaned_data.get("password")
            user = authenticate(request, username=username, password=password)
            if user:
                login(request, user)
                messages.success(request, "Login successful!")
                return redirect("dashboard_view")
            messages.error(request, "Invalid username or password.")
        else:
            messages.error(request, "Please correct the errors below.")
    else:
        form = CustomAuthenticationForm()
    return render(request, "myapp/login.html", {"form": form})


# --------------------------- CUSTOM LOGOUT ------------------------------------
class CustomLogoutView(LogoutView):
    next_page = "login"

    def dispatch(self, request, *args, **kwargs):
        messages.success(request, "You have logged out successfully!")
        return super().dispatch(request, *args, **kwargs)


# -------------------------------- DASHBOARD -----------------------------------
@never_cache
@login_required
@cache_control(no_cache=True, must_revalidate=True, no_store=True)
def dashboard_view(request):
    total_candidates = Candidates.objects.count()
    pending_candidates = Candidates.objects.filter(candidate_status="Pending").count()
    travelled_candidates = Candidates.objects.filter(candidate_status="Travelled").count()

    context = {
        "total_candidates": total_candidates,
        "pending_candidates": pending_candidates,
        "travelled_candidates": travelled_candidates,
    }
    return render(request, "myapp/dashboard.html", context)


# ------------------------------ ADD CLIENT / FORM -----------------------------
@never_cache
@login_required
@cache_control(no_cache=True, must_revalidate=True, no_store=True)
def add_client(request):
    if request.method == "POST":
        form = CandidateApplicationForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            messages.success(request, "Application submitted successfully!")
            return redirect("add_client")
        messages.error(request, "Please correct the errors below.")
        print(form.errors)
    else:
        form = CandidateApplicationForm()
    return render(request, "myapp/add_client.html", {"form": form})


# ------------------------------ VIEW CLIENTS ----------------------------------
@never_cache
@login_required
@cache_control(no_cache=True, must_revalidate=True, no_store=True)
def view_clients(request):
    context = {
        "candidates": Candidates.objects.all(),
        "jobs": Jobs.objects.all(),
        "countries": Countries.objects.all(),
        "agents": Agents.objects.all(),
    }
    return render(request, "myapp/view_clients.html", context)


# --------------------------- UPDATE CANDIDATES --------------------------------
@never_cache
@login_required
@cache_control(no_cache=True, must_revalidate=True, no_store=True)
def update_candidates(request):
    if request.method == "POST":
        for candidate in Candidates.objects.all():
            candidate.full_name = request.POST.get(f"full_name_{candidate.id}", candidate.full_name)
            candidate.gender = request.POST.get(f"gender_{candidate.id}", candidate.gender)
            candidate.phone_number = request.POST.get(f"phone_number_{candidate.id}", candidate.phone_number)
            dob = request.POST.get(f"date_of_birth_{candidate.id}")
            if dob:
                candidate.date_of_birth = dob
            candidate.passport_number = request.POST.get(f"passport_number_{candidate.id}", candidate.passport_number)

            job_id = request.POST.get(f"job_applied_{candidate.id}")
            if job_id:
                candidate.job_applied_id = job_id
            location_id = request.POST.get(f"job_location_{candidate.id}")
            if location_id:
                candidate.job_location_id = location_id
            agent_id = request.POST.get(f"agent_{candidate.id}")
            if agent_id:
                candidate.referral_info_id = agent_id

            files = request.FILES
            for field in ["profile_picture", "full_photo", "passport_copy", "medical_copy", "interpol"]:
                key = f"{field}_{candidate.id}"
                if key in files:
                    setattr(candidate, field, files[key])
            candidate.save()
        messages.success(request, "All candidates updated successfully.")
        return redirect("view_clients")
    messages.error(request, "Invalid request method.")
    return redirect("view_clients")


# ------------------------------ EXPORT EXCEL ----------------------------------
@never_cache
@login_required
def export_excel(request):
    candidates = Candidates.objects.all().values(
        "full_name", "gender", "phone_number", "passport_number",
        "job_applied__title", "referral_info__full_name"
    )
    df = pd.DataFrame(list(candidates))
    df.rename(columns={
        "full_name": "Full Name",
        "gender": "Gender",
        "phone_number": "Phone Number",
        "passport_number": "Passport Number",
        "job_applied__title": "Job Applied",
        "referral_info__full_name": "Referred By",
    }, inplace=True)

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = 'attachment; filename="candidates.xlsx"'
    with pd.ExcelWriter(response, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Candidates")
    return response


# ------------------------------ IMPORT EXCEL ----------------------------------
@never_cache
@login_required
def import_excel(request):
    if request.method == "POST" and request.FILES.get("excel_file"):
        excel_file = request.FILES["excel_file"]
        try:
            df = pd.read_excel(excel_file)
            df.columns = df.columns.str.strip().str.lower().str.replace(" ", "_")
            if "date_of_birth" in df.columns:
                df["date_of_birth"] = pd.to_datetime(df["date_of_birth"], errors="coerce").dt.date

            created_count, skipped_count = 0, 0
            for _, row in df.iterrows():
                full_name = row.get("full_name")
                passport_number = row.get("passport_number")
                job_title = row.get("job_applied_title")
                referral_name = row.get("referral_full_name")

                if not (full_name and passport_number and job_title and referral_name):
                    skipped_count += 1
                    continue
                if Candidates.objects.filter(passport_number=passport_number).exists():
                    skipped_count += 1
                    continue

                job = Jobs.objects.filter(title__iexact=job_title).first()
                if not job:
                    job = Jobs.objects.create(
                        title=job_title,
                        description="Imported job - no description provided.",
                        location="Not specified",
                        salary=0,
                        closing_date=date.today() + timedelta(days=30),
                        responsibilities="Imported job - responsibilities not specified.",
                        status="open",
                    )
                referral = Agents.objects.filter(full_name__iexact=referral_name).first()
                if not referral:
                    referral = Agents.objects.create(full_name=referral_name)

                Candidates.objects.create(
                    full_name=full_name,
                    gender=row.get("gender"),
                    phone_number=row.get("phone_number"),
                    passport_number=passport_number,
                    date_of_birth=row.get("date_of_birth"),
                    job_applied=job,
                    referral_info=referral,
                )
                created_count += 1

            if created_count:
                messages.success(request, f"{created_count} candidate(s) imported successfully.")
            if skipped_count:
                messages.info(request, f"{skipped_count} row(s) skipped (duplicates/missing).")

        except Exception as e:
            messages.error(request, f"Import failed: {e}")
        return redirect("view_clients")

    messages.error(request, "No file uploaded.")
    return redirect("view_clients")


# ------------------------------ VIEW CANDIDATE --------------------------------
@never_cache
@login_required
@cache_control(no_cache=True, must_revalidate=True, no_store=True)
def view_candidate(request, candidate_id):
    candidate = get_object_or_404(Candidates, id=candidate_id)
    job = candidate.job_applied
    return render(request, "myapp/view_candidate.html", {"c": candidate, "job": job})


# ------------------------------ DOWNLOAD CV PDF --------------------------------
@login_required
def download_cv_pdf(request, candidate_id):
    candidate = get_object_or_404(Candidates, id=candidate_id)
    template_path = "myapp/cv_template.html"  # Use same template as view
    context = {"c": candidate, "today_date": date.today(), "applicant_number": candidate.id}

    template = get_template(template_path)
    html = template.render(context)

    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="CV_{candidate.full_name}.pdf"'

    pisa_status = pisa.CreatePDF(html, dest=response)
    if pisa_status.err:
        return HttpResponse("Error generating PDF.")
    return response


# ------------------------------ DOWNLOAD CV WORD --------------------------------
@login_required
def download_cv_word(request, candidate_id):
    candidate = get_object_or_404(Candidates, id=candidate_id)
    document = Document()
    document.add_heading("CURRICULUM VITAE", 0)
    document.add_paragraph(f"Company: CARBIB")
    document.add_paragraph(f"Date: {date.today()}")
    document.add_paragraph(f"Applicant No: {candidate.id}")

    # Main info table
    table = document.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Job Applied For"
    table.cell(0, 1).text = candidate.job_applied.title
    table.cell(1, 0).text = "Full Name"
    table.cell(1, 1).text = candidate.full_name
    table.cell(2, 0).text = "Gender"
    table.cell(2, 1).text = candidate.gender
    table.cell(3, 0).text = "Phone"
    table.cell(3, 1).text = candidate.phone_number
    table.cell(4, 0).text = "Passport Number"
    table.cell(4, 1).text = candidate.passport_number
    table.cell(5, 0).text = "Date of Birth"
    table.cell(5, 1).text = str(candidate.date_of_birth)

    # Personal details
    document.add_heading("Personal Details", level=1)
    document.add_paragraph(f"Marital Status: {candidate.marital_status or 'N/A'}")
    document.add_paragraph(f"Education Level: {candidate.education_level or 'N/A'}")

    # Work experience
    document.add_heading("Work Experience", level=1)
    document.add_paragraph(candidate.working_experience or "N/A")

    # Skills
    document.add_heading("Skills", level=1)
    document.add_paragraph(candidate.skills or "N/A")

    # Full photo
    if candidate.full_photo:
        document.add_page_break()
        document.add_heading("Full Photo", level=1)
        document.add_picture(candidate.full_photo.path, width=Inches(2.5))

    # Passport copy
    if candidate.passport_copy:
        document.add_page_break()
        document.add_heading("Passport Copy", level=1)
        document.add_picture(candidate.passport_copy.path, width=Inches(2.5))

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="CV_{candidate.full_name}.docx"'
    document.save(response)
    return response
